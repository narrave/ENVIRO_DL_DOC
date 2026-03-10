#!/usr/bin/env python3
"""
Convert USER_GUIDE.md → USER_GUIDE.docx
Professional layout: numbered headings, TOC, table captions, clean styling.
Run:  python md_to_docx.py
"""

import re, os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

INPUT_FILE  = r"c:\Users\Marketing\Desktop\EIPL\ENVIRO_DL_DOC\USER_GUIDE.md"
OUTPUT_FILE = r"c:\Users\Marketing\Desktop\EIPL\ENVIRO_DL_DOC\USER_GUIDE_v5.docx"

# ── Palette ────────────────────────────────────────────────────────────────
C_NAVY   = RGBColor(0x17, 0x37, 0x5E)
C_BLUE   = RGBColor(0x1F, 0x5C, 0x99)
C_SLATE  = RGBColor(0x2E, 0x74, 0xB5)
C_WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
C_DARK   = RGBColor(0x1A, 0x1A, 0x1A)
C_INLINE = RGBColor(0x7B, 0x00, 0x00)
C_NOTE   = RGBColor(0x44, 0x44, 0x44)

HEX_HDR  = "17375E"
HEX_SUB  = "2E74B5"
HEX_ALT  = "D6E4F0"
HEX_CODE = "F5F5F5"

BODY_PT  = 14
CODE_PT  = 11

# ── Section / table counters ───────────────────────────────────────────────
_sec     = [0, 0, 0]   # levels 1-3
_tbl_n   = [0]

def _bump(level):
    """Increment section counter at *level* (1-based), reset deeper ones."""
    _sec[level - 1] += 1
    for i in range(level, len(_sec)):
        _sec[i] = 0
    return ".".join(str(x) for x in _sec[:level] if x > 0)

def _next_tbl():
    _tbl_n[0] += 1
    return _tbl_n[0]

# ── XML helpers ────────────────────────────────────────────────────────────

def _shd(elem, fill):
    s = OxmlElement("w:shd")
    s.set(qn("w:val"),   "clear")
    s.set(qn("w:color"), "auto")
    s.set(qn("w:fill"),  fill)
    elem.append(s)

def cell_bg(cell, fill):
    _shd(cell._tc.get_or_add_tcPr(), fill)

def cell_valign(cell, val="center"):
    v = OxmlElement("w:vAlign")
    v.set(qn("w:val"), val)
    cell._tc.get_or_add_tcPr().append(v)

def para_shd(p, fill):
    _shd(p._p.get_or_add_pPr(), fill)

def para_border(p, sides, color="AAAAAA", sz="4", sp="4"):
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    for side in sides:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    sz)
        el.set(qn("w:space"), sp)
        el.set(qn("w:color"), color)
        pBdr.append(el)
    pPr.append(pBdr)

def hr(doc, color="2E74B5", sz="12"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(6)
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    sz)
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"),  color)
    pBdr.append(bot)
    pPr.append(pBdr)

def add_toc_field(doc):
    """Insert Word's automatic TOC field (user presses F9 to refresh)."""
    p   = doc.add_paragraph()
    run = p.add_run()
    # begin
    fc1 = OxmlElement("w:fldChar"); fc1.set(qn("w:fldCharType"), "begin")
    run._r.append(fc1)
    # instruction
    it = OxmlElement("w:instrText")
    it.set(qn("xml:space"), "preserve")
    it.text = ' TOC \\o "1-3" \\h \\z \\u '
    run._r.append(it)
    # separate
    fc2 = OxmlElement("w:fldChar"); fc2.set(qn("w:fldCharType"), "separate")
    run._r.append(fc2)
    # placeholder text
    run2 = p.add_run("[Right-click → Update Field to build Table of Contents]")
    run2.italic = True
    run2.font.color.rgb = RGBColor(0x70, 0x70, 0x70)
    run2.font.size = Pt(10)
    # end
    run3 = p.add_run()
    fc3 = OxmlElement("w:fldChar"); fc3.set(qn("w:fldCharType"), "end")
    run3._r.append(fc3)

# ── Inline markdown ────────────────────────────────────────────────────────

_INLINE = re.compile(
    r"\*\*(.+?)\*\*"
    r"|\*(.+?)\*"
    r"|`([^`]+)`"
    r"|\[([^\]]+)\]\([^)]*\)"
)

def add_inline(para, text, *, bold=False, italic=False,
               pt=BODY_PT, color=C_DARK, font="Calibri"):
    last = 0
    for m in _INLINE.finditer(text):
        plain = text[last:m.start()]
        if plain:
            r = para.add_run(plain)
            r.bold = bold; r.italic = italic
            r.font.name = font; r.font.size = Pt(pt)
            r.font.color.rgb = color

        g = m.group(0)
        if g.startswith("**"):
            r = para.add_run(m.group(1))
            r.bold = True; r.italic = italic
            r.font.name = font; r.font.size = Pt(pt)
            r.font.color.rgb = color
        elif g.startswith("*"):
            r = para.add_run(m.group(2))
            r.bold = bold; r.italic = True
            r.font.name = font; r.font.size = Pt(pt)
            r.font.color.rgb = color
        elif g.startswith("`"):
            r = para.add_run(m.group(3))
            r.font.name = "Consolas"; r.font.size = Pt(pt - 0.5)
            r.font.color.rgb = C_INLINE
        else:
            r = para.add_run(m.group(4))
            r.bold = bold; r.italic = italic
            r.font.name = font; r.font.size = Pt(pt)
            r.font.color.rgb = color
        last = m.end()

    tail = text[last:]
    if tail:
        r = para.add_run(tail)
        r.bold = bold; r.italic = italic
        r.font.name = font; r.font.size = Pt(pt)
        r.font.color.rgb = color

def strip_md(t):
    t = re.sub(r"\*\*(.+?)\*\*", r"\1", t)
    t = re.sub(r"\*(.+?)\*",     r"\1", t)
    t = re.sub(r"`([^`]+)`",     r"\1", t)
    t = re.sub(r"\[([^\]]+)\]\([^)]*\)", r"\1", t)
    return t

# ── Heading emitter ────────────────────────────────────────────────────────

def emit_heading(doc, raw_text, md_level):
    """md_level: 1=#, 2=##, 3=###, 4=####"""
    text = strip_md(raw_text)

    # Strip any existing leading numeric prefix (e.g. "8.3 " or "3. ")
    text = re.sub(r'^\d+(\.\d+)*\.?\s+', '', text)

    # Map: ## → section 1, ### → section 2, #### → section 3
    # Skip # (document title)
    if md_level == 1:
        return  # title is on the title page already

    word_level = md_level - 1   # 1, 2, or 3
    num_str    = _bump(word_level)
    full_title = f"{num_str}  {text}"

    p = doc.add_paragraph()
    p.style = f"Heading {min(word_level, 3)}"

    if word_level == 1:
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after  = Pt(4)
        para_border(p, ["bottom"], color="2E74B5", sz="8", sp="4")
        r = p.add_run(full_title)
        r.bold = True
        r.font.name = "Calibri"
        r.font.size = Pt(34)
        r.font.color.rgb = C_NAVY

    elif word_level == 2:
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after  = Pt(3)
        r = p.add_run(full_title)
        r.bold = True
        r.font.name = "Calibri"
        r.font.size = Pt(22)
        r.font.color.rgb = C_BLUE

    else:
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after  = Pt(2)
        r = p.add_run(full_title)
        r.bold = True
        r.font.name = "Calibri"
        r.font.size = Pt(18)
        r.font.color.rgb = C_SLATE

# ── Code block emitter ─────────────────────────────────────────────────────

def emit_code(doc, lines):
    if not lines:
        return
    sp = doc.add_paragraph()
    sp.paragraph_format.space_before = Pt(4)
    sp.paragraph_format.space_after  = Pt(0)

    for idx, cline in enumerate(lines):
        p = doc.add_paragraph()
        p.style = doc.styles["No Spacing"]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        p.paragraph_format.left_indent  = Inches(0.3)

        r = p.add_run(cline if cline.strip() else " ")
        r.font.name = "Consolas"
        r.font.size = Pt(CODE_PT)
        r.font.color.rgb = RGBColor(0x1E, 0x1E, 0x1E)

        para_shd(p, HEX_CODE)

        sides = ["left", "right"]
        if idx == 0:               sides.append("top")
        if idx == len(lines) - 1:  sides.append("bottom")
        para_border(p, sides, color="BBBBBB", sz="4", sp="4")

    sp2 = doc.add_paragraph()
    sp2.paragraph_format.space_before = Pt(0)
    sp2.paragraph_format.space_after  = Pt(8)

# ── Table emitter ──────────────────────────────────────────────────────────

def emit_table(doc, raw_rows, caption_text=""):
    parsed = []
    for raw in raw_rows:
        cells = [c.strip() for c in raw.strip().strip("|").split("|")]
        if all(re.fullmatch(r"[-: ]+", c) for c in cells if c):
            continue
        parsed.append(cells)

    if not parsed:
        return

    n_cols = max(len(r) for r in parsed)
    tbl = doc.add_table(rows=len(parsed), cols=n_cols)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    for r_idx, row_data in enumerate(parsed):
        row = tbl.rows[r_idx]
        row.height = Pt(20)
        for c_idx in range(n_cols):
            cell     = row.cells[c_idx]
            raw_text = row_data[c_idx] if c_idx < len(row_data) else ""
            para     = cell.paragraphs[0]
            para.clear()
            para.paragraph_format.space_before = Pt(3)
            para.paragraph_format.space_after  = Pt(3)
            para.paragraph_format.left_indent  = Inches(0.06)
            cell_valign(cell, "center")

            if r_idx == 0:
                cell_bg(cell, HEX_HDR)
                add_inline(para, raw_text, bold=True, color=C_WHITE, pt=13)
            else:
                if r_idx % 2 == 0:
                    cell_bg(cell, HEX_ALT)
                add_inline(para, raw_text, pt=13)

    # Table caption below
    n = _next_tbl()
    cap_str = f"Table {n}" + (f": {caption_text}" if caption_text else "")
    cap = doc.add_paragraph()
    cap.paragraph_format.space_before = Pt(2)
    cap.paragraph_format.space_after  = Pt(10)
    r = cap.add_run(cap_str)
    r.italic = True
    r.font.name = "Calibri"
    r.font.size = Pt(9.5)
    r.font.color.rgb = C_NOTE

# ── Title page ─────────────────────────────────────────────────────────────

def make_title_page(doc):
    # ── Company banner (top, navy) ──────────────────────────────────────────
    banner = doc.add_table(rows=1, cols=1)
    banner.style = "Table Grid"
    bc = banner.rows[0].cells[0]
    cell_bg(bc, HEX_HDR)
    bp = bc.paragraphs[0]
    bp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    bp.paragraph_format.space_before = Pt(8)
    bp.paragraph_format.space_after  = Pt(8)
    bp.paragraph_format.left_indent  = Inches(0.15)
    rb = bp.add_run("EIPL  \u2014  Envirotech Instruments Pvt. Ltd.")
    rb.font.name = "Calibri"; rb.font.size = Pt(12)
    rb.font.color.rgb = RGBColor(0xBD, 0xD7, 0xEE)

    # ── Product title (big, centered, with top space for visual breathing) ──
    pt_p = doc.add_paragraph()
    pt_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pt_p.paragraph_format.space_before = Pt(72)   # ~1 inch push-down
    pt_p.paragraph_format.space_after  = Pt(6)
    r1 = pt_p.add_run("ENVIRO_DL STM32")
    r1.bold = True; r1.font.name = "Calibri"; r1.font.size = Pt(44)
    r1.font.color.rgb = C_NAVY

    ug_p = doc.add_paragraph()
    ug_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ug_p.paragraph_format.space_before = Pt(0)
    ug_p.paragraph_format.space_after  = Pt(10)
    r2 = ug_p.add_run("User Guide")
    r2.font.name = "Calibri"; r2.font.size = Pt(28)
    r2.font.color.rgb = C_BLUE

    # Rule under subtitle
    hr(doc, color=HEX_HDR, sz="16")

    # ── Metadata table ──────────────────────────────────────────────────────
    sp_before_meta = doc.add_paragraph()
    sp_before_meta.paragraph_format.space_before = Pt(0)
    sp_before_meta.paragraph_format.space_after  = Pt(16)

    meta = [
        ("Product",           "ENVIRO_DL Environmental Data Logger"),
        ("Platform",          "STM32F417 (ARM Cortex-M4)"),
        ("Firmware",          "ENVIRO_DL_STM32F"),
        ("Manufacturer",      "EIPL (Envirotech Instruments Pvt. Ltd.)"),
        ("Document Version",  "1.0"),
        ("Date",              "March 2026"),
    ]
    mt = doc.add_table(rows=len(meta), cols=2)
    mt.style = "Table Grid"
    mt.alignment = WD_TABLE_ALIGNMENT.LEFT

    for i, (lbl, val) in enumerate(meta):
        row = mt.rows[i]
        row.height = Pt(24)
        cells = row.cells
        cell_bg(cells[0], HEX_HDR)
        pl = cells[0].paragraphs[0]
        pl.paragraph_format.space_before = Pt(4)
        pl.paragraph_format.space_after  = Pt(4)
        pl.paragraph_format.left_indent  = Inches(0.1)
        rl = pl.add_run(lbl)
        rl.bold = True; rl.font.name = "Calibri"; rl.font.size = Pt(13)
        rl.font.color.rgb = C_WHITE
        cell_valign(cells[0], "center")

        cell_bg(cells[1], "EEF4FB" if i % 2 == 0 else "FFFFFF")
        pv = cells[1].paragraphs[0]
        pv.paragraph_format.space_before = Pt(4)
        pv.paragraph_format.space_after  = Pt(4)
        pv.paragraph_format.left_indent  = Inches(0.15)
        rv = pv.add_run(val)
        rv.font.name = "Calibri"; rv.font.size = Pt(13)
        rv.font.color.rgb = C_DARK
        cell_valign(cells[1], "center")

    # ── Footer note (plain, centered) ───────────────────────────────────────
    fp = doc.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.paragraph_format.space_before = Pt(40)
    fp.paragraph_format.space_after  = Pt(4)
    rf = fp.add_run(
        "Document generated from firmware source code  \u2014  ENVIRO_DL_STM32F, March 2026\n"
        "For technical support, contact SVIOT."
    )
    rf.italic = True; rf.font.name = "Calibri"; rf.font.size = Pt(10)
    rf.font.color.rgb = C_NOTE

    doc.add_page_break()

# ── TOC page ───────────────────────────────────────────────────────────────

def make_toc_page(doc):
    # Heading
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(12)
    r = p.add_run("Table of Contents")
    r.bold = True; r.font.name = "Calibri"; r.font.size = Pt(28)
    r.font.color.rgb = C_NAVY
    para_border(p, ["bottom"], color="2E74B5", sz="8", sp="4")

    add_toc_field(doc)
    doc.add_page_break()

# ── Document style setup ───────────────────────────────────────────────────

def setup_doc(doc):
    # Page margins
    for sec in doc.sections:
        sec.top_margin    = Cm(2.5)
        sec.bottom_margin = Cm(2.5)
        sec.left_margin   = Cm(2.8)
        sec.right_margin  = Cm(2.5)

    # Normal style
    n = doc.styles["Normal"]
    n.font.name = "Calibri"
    n.font.size = Pt(BODY_PT)
    n.font.color.rgb = C_DARK

    # Pre-tune heading styles so Word's heading paragraph styles
    # don't fight our run-level formatting
    for lvl, sz, clr in [(1, 34, C_NAVY), (2, 22, C_BLUE), (3, 18, C_SLATE)]:
        try:
            s = doc.styles[f"Heading {lvl}"]
            s.font.name  = "Calibri"
            s.font.size  = Pt(sz)
            s.font.bold  = True
            s.font.color.rgb = clr
        except KeyError:
            pass

# ── Caption / note helpers ─────────────────────────────────────────────────

def emit_caption(doc, text):
    """Inline bold label before a code block or diagram."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    add_inline(p, text, bold=True, pt=BODY_PT, color=C_DARK)

def emit_note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.4)
    p.paragraph_format.right_indent = Inches(0.2)
    p.paragraph_format.space_after  = Pt(5)
    para_border(p, ["left"], color="2E74B5", sz="10", sp="8")
    para_shd(p, "EEF4FB")
    add_inline(p, text, italic=True, color=C_NOTE, pt=BODY_PT - 0.5)

def emit_para(doc, text):
    if not text.strip():
        return
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.space_before = Pt(0)
    add_inline(p, text)

# ── Main conversion ────────────────────────────────────────────────────────

SKIP_LINES = {
    "# ENVIRO_DL STM32 \u2014 User Guide",
    "**Product:** ENVIRO_DL Environmental Data Logger",
    "**Platform:** STM32F417 (ARM Cortex-M4)",
    "**Firmware:** ENVIRO_DL_STM32F",
    "**Manufacturer:** EIPL (Envirotech Instruments Pvt. Ltd.)",
}

def convert(src, dst):
    _sec[:] = [0, 0, 0]
    _tbl_n[0] = 0

    doc = Document()
    setup_doc(doc)
    make_title_page(doc)
    make_toc_page(doc)

    with open(src, encoding="utf-8") as fh:
        lines = fh.readlines()

    in_code    = False
    code_buf   = []
    in_table   = False
    tbl_buf    = []
    para_buf   = []
    skip_md_toc = False   # True while inside the markdown TOC list block

    def flush_para():
        nonlocal para_buf
        text = " ".join(para_buf).strip()
        para_buf = []
        if text:
            emit_para(doc, text)

    def flush_table():
        nonlocal tbl_buf, in_table
        emit_table(doc, tbl_buf)
        tbl_buf  = []
        in_table = False

    for raw in lines:
        line = raw.rstrip("\n")

        if line.strip() in SKIP_LINES:
            continue

        # ── Skip markdown TOC block (## Table of Contents … ---) ────────
        stripped = line.strip()
        if stripped == "## Table of Contents":
            skip_md_toc = True
            continue
        if skip_md_toc:
            if stripped == "---":
                skip_md_toc = False
                hr(doc)        # keep the separator rule in the doc
            continue

        # ── Code fence ──────────────────────────────────────────────────
        if re.match(r"^\s*```", line):
            if in_code:
                if in_table: flush_table()
                flush_para()
                emit_code(doc, code_buf)
                code_buf = []
                in_code  = False
            else:
                if in_table: flush_table()
                flush_para()
                in_code = True
            continue

        if in_code:
            code_buf.append(line)
            continue

        # ── Table row ────────────────────────────────────────────────────
        if re.match(r"^\s*\|", line):
            if not in_table:
                flush_para()
                in_table = True
            tbl_buf.append(line)
            continue
        elif in_table:
            flush_table()

        # ── Blank line ───────────────────────────────────────────────────
        if not line.strip():
            flush_para()
            continue

        # ── Horizontal rule ──────────────────────────────────────────────
        if re.match(r"^---+\s*$", line.strip()):
            flush_para()
            hr(doc)
            continue

        # ── Headings ─────────────────────────────────────────────────────
        m_h = re.match(r"^(#{1,6})\s+(.*)", line)
        if m_h:
            flush_para()
            emit_heading(doc, m_h.group(2).strip(), len(m_h.group(1)))
            continue

        # ── Unordered list ───────────────────────────────────────────────
        m_ul = re.match(r"^(\s*)[-*]\s+(.*)", line)
        if m_ul:
            flush_para()
            depth = len(m_ul.group(1)) // 2
            style = "List Bullet" if depth == 0 else "List Bullet 2"
            p = doc.add_paragraph(style=style)
            p.paragraph_format.space_after  = Pt(2)
            p.paragraph_format.space_before = Pt(1)
            add_inline(p, m_ul.group(2).strip(), pt=BODY_PT)
            continue

        # ── Ordered list ─────────────────────────────────────────────────
        m_ol = re.match(r"^(\s*)\d+[.)]\s+(.*)", line)
        if m_ol:
            flush_para()
            depth = len(m_ol.group(1)) // 3
            style = "List Number" if depth == 0 else "List Number 2"
            p = doc.add_paragraph(style=style)
            p.paragraph_format.space_after  = Pt(2)
            p.paragraph_format.space_before = Pt(1)
            add_inline(p, m_ol.group(2).strip(), pt=BODY_PT)
            continue

        # ── Blockquote ───────────────────────────────────────────────────
        if line.startswith(">"):
            flush_para()
            text = re.sub(r"\\?\*\s*", "", line.lstrip("> ").strip())
            emit_note(doc, text)
            continue

        # ── Normal paragraph ─────────────────────────────────────────────
        para_buf.append(line)

    # Flush remaining
    flush_para()
    if in_code:  emit_code(doc, code_buf)
    if in_table: flush_table()

    doc.save(dst)
    kb = os.path.getsize(dst) // 1024
    print(f"[OK]  {dst}  ({kb} KB)")

if __name__ == "__main__":
    convert(INPUT_FILE, OUTPUT_FILE)
